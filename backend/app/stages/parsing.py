import csv
import io
import json
from collections.abc import Callable
from datetime import datetime
from pathlib import Path

import frontmatter
import pypdfium2 as pdfium
import structlog
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from pypdf import PdfReader

from app.config import Settings
from app.stages.base import Document

logger = structlog.get_logger()

# 页级进度回调：(已完成页数, 总页数)
PageProgressCallback = Callable[[int, int], None]
# 文件级进度回调：(文件名, 已完成页数, 总页数)
FileProgressCallback = Callable[[str, int, int], None]
# OCR 回调：输入页图 PNG 字节，返回识别文本
OcrCallback = Callable[[bytes], str]


class FileParseError(ValueError):
    """知识库文件解析失败（损坏、加密、无文字层等），消息面向用户可读。"""


class NativeParser:
    """Native 知识源解析器，支持 Markdown / JSON / CSV / TXT / PDF / DOCX / XLSX。"""

    def __init__(
        self,
        data_dir: Path,
        pdf_batch_page_threshold: int = 50,
        pdf_page_batch_size: int = 25,
        ocr_fn: OcrCallback | None = None,
        pdf_ocr_dpi: int = 200,
        pdf_max_pages: int = 2000,
    ):
        self.data_dir = Path(data_dir)
        # 超过阈值页数的 PDF 按批推送解析心跳，小 PDF 不增加事件噪音
        self.pdf_batch_page_threshold = pdf_batch_page_threshold
        self.pdf_page_batch_size = pdf_page_batch_size
        # 单份 PDF 页数上限，超限直接拒绝
        self.pdf_max_pages = pdf_max_pages
        # 无文字层 PDF 页的回退 OCR；为 None 时扫描件按失败处理
        self.ocr_fn = ocr_fn
        self.pdf_ocr_dpi = pdf_ocr_dpi
        # 全量解析中被跳过文件的失败原因（"文件名: 原因"），供任务日志展示
        self.parse_errors: list[str] = []

    def parse(self, on_progress: FileProgressCallback | None = None) -> list[Document]:
        self.parse_errors = []
        documents: list[Document] = []
        for subdir in ["help_docs", "changelog", "tickets", "faq"]:
            path = self.data_dir / subdir
            if not path.exists():
                continue
            if subdir == "help_docs":
                documents.extend(self._parse_help_docs(path))
            elif subdir == "changelog":
                documents.extend(self._parse_changelog(path))
            elif subdir == "tickets":
                documents.extend(self._parse_tickets(path))
            elif subdir == "faq":
                documents.extend(self._parse_faq(path))

        # 管理员后台上传的知识库文件统一放在 data/kb 下
        kb_path = self.data_dir / "kb"
        if kb_path.exists():
            documents.extend(self._parse_kb(kb_path, on_progress=on_progress))

        return documents

    @staticmethod
    def _now() -> datetime:
        return datetime.utcnow()

    def _parse_help_docs(self, path: Path) -> list[Document]:
        docs: list[Document] = []
        for file_path in sorted(path.rglob("*.md")):
            post = frontmatter.load(file_path)
            title = post.get("title") or file_path.stem
            updated_at = self._parse_datetime(post.get("updated_at")) or self._now()
            docs.append(
                Document(
                    content=post.content.strip(),
                    source_type="help_doc",
                    title=str(title),
                    updated_at=updated_at,
                    source_id=str(file_path.relative_to(self.data_dir)),
                )
            )
        return docs

    def _parse_changelog(self, path: Path) -> list[Document]:
        docs: list[Document] = []
        for file_path in sorted(path.rglob("*.json")):
            data = json.loads(file_path.read_text(encoding="utf-8"))
            entries = data if isinstance(data, list) else [data]
            for idx, entry in enumerate(entries):
                docs.append(
                    Document(
                        content=f"{entry.get('title', '')}\n{entry.get('content', '')}".strip(),
                        source_type="changelog",
                        title=str(entry.get("title", file_path.stem)),
                        updated_at=self._parse_datetime(entry.get("date")) or self._now(),
                        source_id=f"{file_path.relative_to(self.data_dir)}#{idx}",
                    )
                )
        return docs

    def _parse_tickets(self, path: Path) -> list[Document]:
        docs: list[Document] = []
        for file_path in sorted(path.rglob("*.csv")):
            with file_path.open("r", encoding="utf-8", newline="") as f:
                reader = csv.DictReader(f)
                for idx, row in enumerate(reader):
                    content = "\n".join(f"{k}: {v}" for k, v in row.items())
                    docs.append(
                        Document(
                            content=content,
                            source_type="ticket",
                            title=f"工单 #{row.get('ticket_id', idx)}",
                            updated_at=self._parse_datetime(row.get("updated_at")) or self._now(),
                            source_id=f"{file_path.relative_to(self.data_dir)}#row_{idx}",
                        )
                    )
        return docs

    def _parse_faq(self, path: Path) -> list[Document]:
        docs: list[Document] = []
        for file_path in sorted(path.rglob("*.json")):
            data = json.loads(file_path.read_text(encoding="utf-8"))
            entries = data if isinstance(data, list) else [data]
            for idx, entry in enumerate(entries):
                q = entry.get("question", "")
                a = entry.get("answer", "")
                docs.append(
                    Document(
                        content=f"Q: {q}\nA: {a}".strip(),
                        source_type="faq",
                        title=q or file_path.stem,
                        updated_at=self._parse_datetime(entry.get("updated_at")) or self._now(),
                        source_id=f"{file_path.relative_to(self.data_dir)}#{idx}",
                    )
                )
        return docs

    SUPPORTED_SUFFIXES = {".md", ".json", ".csv", ".txt", ".pdf", ".docx", ".xlsx"}
    # 新格式（二进制）解析失败必须给出可读原因，旧格式维持"失败返回空"的历史行为
    _RAISE_ON_FAILURE_SUFFIXES = {".pdf", ".docx", ".xlsx"}

    def parse_file(
        self,
        file_path: Path,
        on_progress: PageProgressCallback | None = None,
    ) -> list[Document]:
        """解析单个知识库文件，返回一个或多个 Document。

        不支持的后缀返回空列表；PDF/DOCX/XLSX 解析失败抛出 FileParseError（可读原因）。
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_SUFFIXES:
            return []

        try:
            if suffix == ".md":
                post = frontmatter.load(file_path)
                content = post.content.strip()
                title = post.get("title") or file_path.stem
            elif suffix == ".json":
                data = json.loads(file_path.read_text(encoding="utf-8"))
                entries = data if isinstance(data, list) else [data]
                return [
                    Document(
                        content=json.dumps(entry, ensure_ascii=False),
                        source_type="kb_doc",
                        title=file_path.stem,
                        updated_at=self._now(),
                        source_id=str(file_path.relative_to(self.data_dir)),
                    )
                    for entry in entries
                ]
            elif suffix == ".csv":
                rows: list[str] = []
                with file_path.open("r", encoding="utf-8", newline="") as f:
                    reader = csv.DictReader(f)
                    for idx, row in enumerate(reader):
                        rows.append("\n".join(f"{k}: {v}" for k, v in row.items()))
                content = "\n\n".join(rows)
                title = file_path.stem
            elif suffix == ".txt":
                content = file_path.read_text(encoding="utf-8").strip()
                title = file_path.stem
            elif suffix == ".pdf":
                content = self._parse_pdf(file_path, on_progress)
                title = file_path.stem
            elif suffix == ".docx":
                content = self._parse_docx(file_path)
                title = file_path.stem
            else:  # .xlsx
                content = self._parse_xlsx(file_path)
                title = file_path.stem
        except FileParseError:
            raise
        except Exception as exc:
            if suffix in self._RAISE_ON_FAILURE_SUFFIXES:
                raise FileParseError(f"{file_path.name} 解析失败：{exc}") from exc
            return []

        relative = file_path.relative_to(self.data_dir)
        source_type = "kb_doc"
        parts = relative.parts
        if len(parts) > 2:
            source_type = f"kb_{parts[1]}"

        return [
            Document(
                content=content,
                source_type=source_type,
                title=str(title),
                updated_at=self._now(),
                source_id=str(relative),
            )
        ]

    def _parse_pdf(
        self,
        file_path: Path,
        on_progress: PageProgressCallback | None = None,
    ) -> str:
        """按页提取 PDF 文本并带页码标记；无文字层页回退 OCR；大 PDF 按批推送进度心跳。"""
        try:
            reader = PdfReader(str(file_path))
        except Exception as exc:
            raise FileParseError(f"PDF 文件损坏或无法读取：{exc}") from exc
        if reader.is_encrypted:
            raise FileParseError("PDF 已加密，无法解析")

        total_pages = len(reader.pages)
        if total_pages > self.pdf_max_pages:
            raise FileParseError(
                f"PDF 共 {total_pages} 页，超过 {self.pdf_max_pages} 页上限，请拆分后上传"
            )
        heartbeat = on_progress if total_pages > self.pdf_batch_page_threshold else None
        page_texts: list[str] = []
        ocr_attempted = False
        pdfium_doc = None
        try:
            for index, page in enumerate(reader.pages, start=1):
                try:
                    text = (page.extract_text() or "").strip()
                except Exception as exc:
                    raise FileParseError(f"PDF 第 {index} 页解析失败：{exc}") from exc

                used_ocr = False
                if not text and self.ocr_fn is not None:
                    ocr_attempted = True
                    if pdfium_doc is None:
                        pdfium_doc = pdfium.PdfDocument(str(file_path))
                    try:
                        text = self._ocr_page(pdfium_doc, index - 1)
                        used_ocr = bool(text)
                    except Exception as exc:
                        # 单页 OCR 失败只跳过该页，不拖垮整份文档
                        logger.warning(
                            "pdf_page_ocr_failed", path=str(file_path), page=index, error=str(exc)
                        )

                if text:
                    page_texts.append(f"[第 {index} 页]\n{text}")
                # OCR 页耗时长，逐页心跳；普通大 PDF 仍按批心跳
                if used_ocr and on_progress:
                    on_progress(index, total_pages)
                elif heartbeat and index % self.pdf_page_batch_size == 0:
                    heartbeat(index, total_pages)
        finally:
            if pdfium_doc is not None:
                pdfium_doc.close()

        content = "\n\n".join(page_texts).strip()
        if not content:
            if ocr_attempted:
                raise FileParseError("该 PDF 经 OCR 仍未提取到文本内容")
            if self.ocr_fn is None:
                raise FileParseError("该 PDF 无文字层（可能是扫描件），且 OCR 未启用，无法解析")
            raise FileParseError("PDF 无可用文本内容")
        return content

    def _ocr_page(self, pdfium_doc, page_index: int) -> str:
        """栅格化单页为 PNG 并调用 OCR，返回识别文本。"""
        page = pdfium_doc[page_index]
        bitmap = page.render(scale=self.pdf_ocr_dpi / 72)
        buffer = io.BytesIO()
        bitmap.to_pil().save(buffer, format="PNG")
        return (self.ocr_fn(buffer.getvalue()) or "").strip()

    def _parse_docx(self, file_path: Path) -> str:
        """按文档顺序提取段落与表格（表格按行序列化为 "单元格 | 单元格"）。"""
        document = DocxDocument(str(file_path))
        parts: list[str] = []
        for child in document.element.body.iterchildren():
            if child.tag == qn("w:p"):
                text = Paragraph(child, document).text.strip()
                if text:
                    parts.append(text)
            elif child.tag == qn("w:tbl"):
                for row in Table(child, document).rows:
                    line = " | ".join(cell.text.strip() for cell in row.cells)
                    if line.strip(" |"):
                        parts.append(line)
        content = "\n".join(parts).strip()
        if not content:
            raise FileParseError("Word 文档无可用文本内容")
        return content

    def _parse_xlsx(self, file_path: Path) -> str:
        """逐 sheet 提取，首行作表头，数据行沿用 CSV 路径的 "列名: 值" 格式。"""
        workbook = load_workbook(str(file_path), read_only=True, data_only=True)
        try:
            sheet_blocks: list[str] = []
            for sheet in workbook.worksheets:
                row_iter = sheet.iter_rows(values_only=True)
                header_row = next(row_iter, None)
                if header_row is None:
                    continue
                headers = [str(cell).strip() if cell is not None else "" for cell in header_row]
                lines: list[str] = []
                for row in row_iter:
                    pairs: list[str] = []
                    for idx, value in enumerate(row):
                        if value is None or str(value).strip() == "":
                            continue
                        key = headers[idx] if idx < len(headers) and headers[idx] else f"列{idx + 1}"
                        pairs.append(f"{key}: {value}")
                    if pairs:
                        lines.append("\n".join(pairs))
                if lines:
                    sheet_blocks.append(f"[Sheet: {sheet.title}]\n" + "\n\n".join(lines))
        finally:
            workbook.close()

        content = "\n\n".join(sheet_blocks).strip()
        if not content:
            raise FileParseError("Excel 文件无可用数据")
        return content

    def _parse_kb(
        self,
        kb_root: Path,
        on_progress: FileProgressCallback | None = None,
    ) -> list[Document]:
        """解析管理员后台上传的知识库文件（data/kb 下任意子目录）。

        单文件解析失败只记录并跳过，不拖垮整批重建。
        """
        docs: list[Document] = []
        for file_path in sorted(kb_root.rglob("*")):
            if file_path.is_dir():
                continue
            try:
                file_progress = None
                if on_progress is not None:

                    def file_progress(done: int, total: int, name: str = file_path.name) -> None:
                        on_progress(name, done, total)

                docs.extend(self.parse_file(file_path, on_progress=file_progress))
            except FileParseError as exc:
                logger.warning("kb_file_parse_failed", path=str(file_path), error=str(exc))
                self.parse_errors.append(f"{file_path.name}: {exc}")
        return docs

    def _parse_datetime(self, value) -> datetime | None:
        if not value:
            return None
        if isinstance(value, datetime):
            return value
        if isinstance(value, str):
            for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S"):
                try:
                    return datetime.strptime(value, fmt)
                except ValueError:
                    continue
        return None


def build_parser(settings: Settings, data_dir: Path, model_client=None, settings_service=None):
    """按运行配置构造解析器。全量重建与单文件索引必须共用同一选择逻辑。

    解析器与 OCR/PDF 参数走运行期配置（DB → .env → 默认）；
    提供 model_client 且启用 OCR 时，无文字层 PDF 页回退到视觉模型识别。
    """
    from app.services.settings_service import SettingsService

    svc = settings_service or SettingsService()
    if svc.get_runtime_value("parser") == "llamaindex":
        from app.stages.adapters.li_parsing import LlamaIndexParserAdapter

        return LlamaIndexParserAdapter(data_dir)
    ocr_enabled = svc.get_runtime_value("ocr_enabled")
    ocr_fn = model_client.ocr_image if (model_client is not None and ocr_enabled) else None
    return NativeParser(
        data_dir,
        pdf_batch_page_threshold=svc.get_runtime_value("pdf_batch_page_threshold"),
        pdf_page_batch_size=svc.get_runtime_value("pdf_page_batch_size"),
        ocr_fn=ocr_fn,
        pdf_ocr_dpi=svc.get_runtime_value("pdf_ocr_dpi"),
        pdf_max_pages=svc.get_runtime_value("pdf_max_pages"),
    )
