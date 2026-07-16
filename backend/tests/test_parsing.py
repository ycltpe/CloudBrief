
import pytest

from app.stages.parsing import FileParseError, NativeParser


def test_parse_file_md(tmp_path):
    data_dir = tmp_path / "data"
    kb_dir = data_dir / "kb" / "dir_1"
    kb_dir.mkdir(parents=True)
    file_path = kb_dir / "test.md"
    file_path.write_text("---\ntitle: Hello\n---\n# World\n\ncontent", encoding="utf-8")

    parser = NativeParser(data_dir)
    docs = parser.parse_file(file_path)

    assert len(docs) == 1
    assert docs[0].title == "Hello"
    assert "World" in docs[0].content
    assert docs[0].source_id == "kb/dir_1/test.md"


def test_parse_file_txt(tmp_path):
    data_dir = tmp_path / "data"
    kb_dir = data_dir / "kb"
    kb_dir.mkdir(parents=True)
    file_path = kb_dir / "note.txt"
    file_path.write_text("plain text", encoding="utf-8")

    parser = NativeParser(data_dir)
    docs = parser.parse_file(file_path)

    assert len(docs) == 1
    assert docs[0].content == "plain text"
    assert docs[0].source_type == "kb_doc"


def test_parse_file_unsupported(tmp_path):
    data_dir = tmp_path / "data"
    file_path = data_dir / "x.doc"
    file_path.parent.mkdir(parents=True)
    file_path.write_text("x")

    parser = NativeParser(data_dir)
    assert parser.parse_file(file_path) == []


def _write_text_pdf(path, pages: list[str]) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=letter)
    for text in pages:
        pdf.drawString(72, 720, text)
        pdf.showPage()
    pdf.save()


def test_parse_file_pdf(tmp_path):
    data_dir = tmp_path / "data"
    kb_dir = data_dir / "kb" / "dir_1"
    kb_dir.mkdir(parents=True)
    file_path = kb_dir / "manual.pdf"
    _write_text_pdf(file_path, ["CloudBrief handbook page one", "page two content"])

    parser = NativeParser(data_dir)
    docs = parser.parse_file(file_path)

    assert len(docs) == 1
    assert docs[0].title == "manual"
    assert "[第 1 页]" in docs[0].content
    assert "CloudBrief handbook page one" in docs[0].content
    assert "[第 2 页]" in docs[0].content
    assert docs[0].source_id == "kb/dir_1/manual.pdf"


def test_parse_file_pdf_progress_only_for_large_pdf(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)

    small_pdf = data_dir / "small.pdf"
    _write_text_pdf(small_pdf, ["p1", "p2", "p3"])
    parser = NativeParser(data_dir)
    calls: list[tuple[int, int]] = []
    parser.parse_file(small_pdf, on_progress=lambda d, t: calls.append((d, t)))
    assert calls == []  # 小 PDF 不产生心跳噪音

    large_pdf = data_dir / "large.pdf"
    _write_text_pdf(large_pdf, [f"page {i}" for i in range(5)])
    parser = NativeParser(data_dir, pdf_batch_page_threshold=2, pdf_page_batch_size=2)
    parser.parse_file(large_pdf, on_progress=lambda d, t: calls.append((d, t)))
    assert calls == [(2, 5), (4, 5)]


def test_parse_file_scanned_pdf_raises(tmp_path):
    from pypdf import PdfWriter

    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "scanned.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with file_path.open("wb") as f:
        writer.write(f)

    parser = NativeParser(data_dir)
    with pytest.raises(FileParseError, match="无文字层"):
        parser.parse_file(file_path)


def test_parse_file_corrupt_pdf_raises(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "broken.pdf"
    file_path.write_bytes(b"this is not a pdf")

    parser = NativeParser(data_dir)
    with pytest.raises(FileParseError, match="无法读取"):
        parser.parse_file(file_path)


def test_parse_file_docx_keeps_order_and_tables(tmp_path):
    from docx import Document as DocxDocument

    data_dir = tmp_path / "data"
    kb_dir = data_dir / "kb"
    kb_dir.mkdir(parents=True)
    file_path = kb_dir / "sop.docx"

    document = DocxDocument()
    document.add_paragraph("第一段：目的")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "步骤"
    table.rows[0].cells[1].text = "说明"
    document.add_paragraph("第二段：收尾")
    document.save(str(file_path))

    parser = NativeParser(data_dir)
    docs = parser.parse_file(file_path)

    assert len(docs) == 1
    content = docs[0].content
    assert "步骤 | 说明" in content
    assert content.index("第一段") < content.index("步骤 | 说明") < content.index("第二段")


def test_parse_file_docx_empty_raises(tmp_path):
    from docx import Document as DocxDocument

    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "empty.docx"
    DocxDocument().save(str(file_path))

    parser = NativeParser(data_dir)
    with pytest.raises(FileParseError, match="无可用文本内容"):
        parser.parse_file(file_path)


def test_parse_file_xlsx_multi_sheet(tmp_path):
    from openpyxl import Workbook

    data_dir = tmp_path / "data"
    kb_dir = data_dir / "kb"
    kb_dir.mkdir(parents=True)
    file_path = kb_dir / "faq.xlsx"

    workbook = Workbook()
    ws1 = workbook.active
    ws1.title = "FAQ"
    ws1.append(["问题", "答案"])
    ws1.append(["如何重置密码", "在设置页操作"])
    ws2 = workbook.create_sheet("联系方式")
    ws2.append(["部门", "电话"])
    ws2.append(["客服", "10086"])
    workbook.save(str(file_path))

    parser = NativeParser(data_dir)
    docs = parser.parse_file(file_path)

    assert len(docs) == 1
    content = docs[0].content
    assert "[Sheet: FAQ]" in content
    assert "问题: 如何重置密码" in content
    assert "答案: 在设置页操作" in content
    assert "[Sheet: 联系方式]" in content
    assert "部门: 客服" in content


def test_parse_file_xlsx_empty_raises(tmp_path):
    from openpyxl import Workbook

    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "empty.xlsx"
    Workbook().save(str(file_path))

    parser = NativeParser(data_dir)
    with pytest.raises(FileParseError, match="无可用数据"):
        parser.parse_file(file_path)


def test_parse_kb_skips_failed_file_without_aborting(tmp_path):
    from pypdf import PdfWriter

    data_dir = tmp_path / "data"
    kb_dir = data_dir / "kb" / "dir_1"
    kb_dir.mkdir(parents=True)
    (kb_dir / "good.txt").write_text("有效内容", encoding="utf-8")
    scanned = kb_dir / "scanned.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with scanned.open("wb") as f:
        writer.write(f)

    parser = NativeParser(data_dir)
    docs = parser.parse()

    assert len(docs) == 1
    assert docs[0].content == "有效内容"
    assert len(parser.parse_errors) == 1
    assert "scanned.pdf" in parser.parse_errors[0]
    assert "无文字层" in parser.parse_errors[0]


def _write_blank_pdf(path, pages: int) -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=612, height=792)
    with path.open("wb") as f:
        writer.write(f)


def _write_mixed_pdf(path) -> None:
    """两页原生文本夹一页空白（无文字层）的混合 PDF。"""
    from pypdf import PdfReader, PdfWriter

    text_pdf = path.parent / "_text_pages.pdf"
    _write_text_pdf(text_pdf, ["native page one", "native page two"])
    reader = PdfReader(str(text_pdf))
    writer = PdfWriter()
    writer.add_page(reader.pages[0])
    writer.add_blank_page(width=612, height=792)
    writer.add_page(reader.pages[1])
    with path.open("wb") as f:
        writer.write(f)
    text_pdf.unlink()


def test_parse_pdf_scanned_uses_ocr_with_progress(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "scanned.pdf"
    _write_blank_pdf(file_path, 2)

    ocr_calls: list[int] = []

    def fake_ocr(png_bytes: bytes) -> str:
        ocr_calls.append(len(png_bytes))
        return "OCR识别出的文字"

    progress: list[tuple[int, int]] = []
    parser = NativeParser(data_dir, ocr_fn=fake_ocr)
    docs = parser.parse_file(file_path, on_progress=lambda d, t: progress.append((d, t)))

    assert len(docs) == 1
    assert "[第 1 页]\nOCR识别出的文字" in docs[0].content
    assert "[第 2 页]\nOCR识别出的文字" in docs[0].content
    assert len(ocr_calls) == 2 and all(size > 0 for size in ocr_calls)
    # OCR 页即使低于大 PDF 阈值也逐页心跳
    assert progress == [(1, 2), (2, 2)]


def test_parse_pdf_mixed_ocr_only_blank_pages(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "mixed.pdf"
    _write_mixed_pdf(file_path)

    ocr_calls: list[int] = []

    def fake_ocr(png_bytes: bytes) -> str:
        ocr_calls.append(len(png_bytes))
        return "空白页识别结果"

    parser = NativeParser(data_dir, ocr_fn=fake_ocr)
    docs = parser.parse_file(file_path)

    assert len(docs) == 1
    assert "native page one" in docs[0].content
    assert "native page two" in docs[0].content
    assert "[第 2 页]\n空白页识别结果" in docs[0].content
    # 只有无文字层的第 2 页走了 OCR
    assert len(ocr_calls) == 1


def test_parse_pdf_ocr_page_failure_isolated(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "scanned.pdf"
    _write_blank_pdf(file_path, 3)

    call_count = {"n": 0}

    def flaky_ocr(png_bytes: bytes) -> str:
        call_count["n"] += 1
        if call_count["n"] == 2:
            raise RuntimeError("VL 服务超时")
        return f"第{call_count['n']}页文字"

    parser = NativeParser(data_dir, ocr_fn=flaky_ocr)
    docs = parser.parse_file(file_path)

    assert len(docs) == 1
    assert "第1页文字" in docs[0].content
    assert "第3页文字" in docs[0].content
    assert "[第 2 页]" not in docs[0].content


def test_parse_pdf_ocr_empty_result_raises(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "scanned.pdf"
    _write_blank_pdf(file_path, 1)

    parser = NativeParser(data_dir, ocr_fn=lambda png: "")
    with pytest.raises(FileParseError, match="OCR 仍未提取到文本"):
        parser.parse_file(file_path)


def test_parse_pdf_over_max_pages_rejected(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir(parents=True)
    file_path = data_dir / "huge.pdf"
    _write_text_pdf(file_path, [f"page {i}" for i in range(5)])

    parser = NativeParser(data_dir, pdf_max_pages=3)
    with pytest.raises(FileParseError, match="上限"):
        parser.parse_file(file_path)


def test_build_parser_wires_ocr_fn(tmp_path, monkeypatch):
    from unittest.mock import MagicMock

    from app.config import get_settings
    from app.stages.parsing import build_parser

    settings = get_settings()
    client = MagicMock()

    def _fake_svc(ocr_enabled: bool):
        svc = MagicMock()
        values = {
            "parser": "native",
            "ocr_enabled": ocr_enabled,
            "pdf_batch_page_threshold": 50,
            "pdf_page_batch_size": 25,
            "pdf_ocr_dpi": 200,
            "pdf_max_pages": 2000,
        }
        svc.get_runtime_value.side_effect = lambda key: values.get(key)
        return svc

    parser = build_parser(settings, tmp_path, model_client=client, settings_service=_fake_svc(True))
    assert isinstance(parser, NativeParser)
    assert parser.ocr_fn == client.ocr_image

    parser = build_parser(settings, tmp_path, model_client=client, settings_service=_fake_svc(False))
    assert parser.ocr_fn is None
